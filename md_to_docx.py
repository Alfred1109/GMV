#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def md_to_docx(md_file, docx_file):
    # Check if input file exists
    if not os.path.isfile(md_file):
        raise FileNotFoundError(f"Input file not found: {md_file}")
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Initialize document
    doc = Document()
    
    # Set document style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add styles if they don't exist
    def get_or_add_style(style_name, style_type):
        if style_name in doc.styles:
            return doc.styles[style_name]
        else:
            return doc.styles.add_style(style_name, style_type)
    
    # Add a title style
    title_style = get_or_add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    
    # Add heading styles
    for i in range(1, 5):
        heading_style = get_or_add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.bold = True
        heading_font.size = Pt(16 - (i * 2))  # Decrease size for deeper headings
    
    # Split markdown by lines
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle headings
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 2')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Heading 3')
        elif line.startswith('##### '):
            doc.add_paragraph(line[6:], style='Heading 4')
            
        # Handle images
        elif line.startswith('![') and ']' in line and '(' in line and ')' in line:
            # Extract alt text and url
            alt_start = line.find('[') + 1
            alt_end = line.find(']')
            alt_text = line[alt_start:alt_end]
            
            url_start = line.find('(') + 1
            url_end = line.find(')')
            url = line[url_start:url_end]
            
            # Add image placeholder
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片占位符: {alt_text}]")
            run.italic = True
            
            # Add image URL as a note
            p = doc.add_paragraph(f"图片路径: {url}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].italic = True
            
            # Check for image caption (usually in next line with * or _)
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('*') and lines[i + 1].strip().endswith('*'):
                caption = lines[i + 1].strip('* ')
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                i += 1  # Skip the caption line in next iteration
            
        # Handle bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Odd indices are bold
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # Handle lists
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        # Handle empty lines
        elif not line:
            doc.add_paragraph()
            
        # Regular text
        else:
            doc.add_paragraph(line)
            
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"成功将 {md_file} 转换为 {docx_file}")

if __name__ == "__main__":
    input_file = "GeniusMed Vault医学科研专病库平台功能模块与产品特点设计方案.md"
    output_file = "GeniusMed Vault医学科研专病库平台功能模块与产品特点设计方案.docx"
    
    md_to_docx(input_file, output_file) 