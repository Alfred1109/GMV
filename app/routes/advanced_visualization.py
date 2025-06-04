"""
高级可视化模块

提供交互式数据可视化和自动报告生成功能
"""

from flask import Blueprint, jsonify, request, current_app, render_template, send_file
from flask_login import login_required, current_user
import json
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from datetime import datetime
import os
from werkzeug.utils import secure_filename
import tempfile
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas

# 添加PDF报告生成相关依赖
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# 添加DOCX报告生成相关依赖
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# 创建Blueprint
advanced_viz_bp = Blueprint('advanced_visualization', __name__)

# 高级图表配置API
@advanced_viz_bp.route('/api/visualization/advanced_chart', methods=['POST', 'OPTIONS'])
@login_required
def advanced_chart_config():
    """高级图表配置API
    
    允许用户自定义图表的各种参数，如颜色、标题、图例等
    
    请求体格式:
    {
        "dataset_id": 1,                     // 数据集ID
        "chart_type": "bar",                 // 图表类型：bar, line, scatter, pie, heatmap, box
        "x_variable": "age",                 // X轴变量
        "y_variable": "bmi",                 // Y轴变量（可选，取决于图表类型）
        "group_variable": "gender",          // 分组变量（可选）
        "config": {                          // 图表配置
            "title": "BMI分布图",            // 图表标题
            "x_label": "年龄",               // X轴标签
            "y_label": "BMI指数",            // Y轴标签
            "color_palette": "viridis",      // 颜色方案
            "figure_size": [10, 6],          // 图表尺寸 [宽, 高]
            "dpi": 100,                      // 分辨率
            "grid": true,                    // 是否显示网格
            "legend": true,                  // 是否显示图例
            "legend_position": "best",       // 图例位置
            "style": "whitegrid",            // 图表样式
            "font_family": "SimHei",         // 字体
            "font_size": 12,                 // 字体大小
            "rotation": 45,                  // X轴标签旋转角度
            "annotations": [                 // 注释（可选）
                {
                    "text": "异常值",
                    "x": 30,
                    "y": 35,
                    "color": "red"
                }
            ]
        }
    }
    
    Returns:
        包含图表数据的JSON响应
    """
    if request.method == 'OPTIONS':
        response = current_app.make_default_options_response()
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST')
        return response
        
    try:
        # 获取请求数据
        request_data = request.get_json()
        if not request_data:
            return jsonify({
                'success': False,
                'message': '请求数据为空'
            }), 400
        
        # 提取参数
        dataset_id = request_data.get('dataset_id')
        chart_type = request_data.get('chart_type')
        x_variable = request_data.get('x_variable')
        y_variable = request_data.get('y_variable')
        group_variable = request_data.get('group_variable')
        config = request_data.get('config', {})
        
        # 验证必要参数
        if not dataset_id:
            return jsonify({
                'success': False,
                'message': '未指定数据集ID'
            }), 400
        
        if not chart_type:
            return jsonify({
                'success': False,
                'message': '未指定图表类型'
            }), 400
        
        if not x_variable:
            return jsonify({
                'success': False,
                'message': '未指定X轴变量'
            }), 400
        
        # 根据图表类型验证其他必要参数
        if chart_type in ['line', 'scatter', 'bar'] and not y_variable:
            return jsonify({
                'success': False,
                'message': f'{chart_type}图表需要指定Y轴变量'
            }), 400
        
        # 检查数据集是否存在，并验证访问权限
        from app.models.dataset import DataSet
        from app.models.dataset_entry import DatasetEntry
        
        dataset = DataSet.query.get(dataset_id)
        if not dataset:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})不存在'
            }), 404
        
        # 检查用户是否有权限访问此数据集
        if dataset.privacy_level and dataset.privacy_level != 'public':
            if dataset.created_by != current_user.id and current_user.role != 'admin':
                if not dataset.is_shared_with(current_user.id):
                    return jsonify({
                        'success': False, 
                        'message': f'没有权限访问数据集(ID={dataset_id})'
                    }), 403
        
        # 获取数据集条目
        entries = DatasetEntry.query.filter_by(dataset_id=dataset_id).all()
        if not entries:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})没有数据条目'
            }), 404
        
        # 提取数据
        data = {}
        variables = [x_variable]
        if y_variable:
            variables.append(y_variable)
        if group_variable:
            variables.append(group_variable)
        
        for var in variables:
            data[var] = []
        
        # 从所有条目中提取所需变量的值
        for entry in entries:
            try:
                entry_data = json.loads(entry.data)
                
                for var in variables:
                    if var in entry_data:
                        value = entry_data[var]
                        # 尝试将字符串转换为数值
                        if isinstance(value, str):
                            try:
                                if '.' in value:
                                    value = float(value)
                                else:
                                    value = int(value)
                            except (ValueError, TypeError):
                                pass
                        data[var].append(value)
                    else:
                        data[var].append(None)  # 缺失值
            except Exception as e:
                current_app.logger.error(f"处理条目数据时出错: {str(e)}")
                continue
        
        # 转换为Pandas DataFrame
        df = pd.DataFrame(data)
        
        # 去除包含缺失值的行
        df = df.dropna()
        
        if len(df) == 0:
            return jsonify({
                'success': False,
                'message': '处理后的数据为空，无法生成图表'
            }), 400
        
        # 应用图表配置
        plt.figure(figsize=tuple(config.get('figure_size', [10, 6])))
        sns.set_style(config.get('style', 'whitegrid'))
        
        # 设置字体
        plt.rcParams['font.family'] = config.get('font_family', 'SimHei')
        plt.rcParams['font.size'] = config.get('font_size', 12)
        
        # 生成图表
        if chart_type == 'bar':
            if group_variable:
                ax = sns.barplot(x=x_variable, y=y_variable, hue=group_variable, data=df, palette=config.get('color_palette', 'viridis'))
            else:
                ax = sns.barplot(x=x_variable, y=y_variable, data=df, palette=config.get('color_palette', 'viridis'))
        
        elif chart_type == 'line':
            if group_variable:
                ax = sns.lineplot(x=x_variable, y=y_variable, hue=group_variable, data=df, palette=config.get('color_palette', 'viridis'))
            else:
                ax = sns.lineplot(x=x_variable, y=y_variable, data=df, palette=config.get('color_palette', 'viridis'))
        
        elif chart_type == 'scatter':
            if group_variable:
                ax = sns.scatterplot(x=x_variable, y=y_variable, hue=group_variable, data=df, palette=config.get('color_palette', 'viridis'))
            else:
                ax = sns.scatterplot(x=x_variable, y=y_variable, data=df, palette=config.get('color_palette', 'viridis'))
        
        elif chart_type == 'pie':
            # 饼图需要特殊处理
            plt.pie(df[y_variable], labels=df[x_variable], autopct='%1.1f%%')
            ax = plt.gca()
        
        elif chart_type == 'heatmap':
            # 热图需要特殊处理，通常用于相关性矩阵
            corr_matrix = df.corr()
            ax = sns.heatmap(corr_matrix, annot=True, cmap=config.get('color_palette', 'viridis'))
        
        elif chart_type == 'box':
            if group_variable:
                ax = sns.boxplot(x=x_variable, y=y_variable, hue=group_variable, data=df, palette=config.get('color_palette', 'viridis'))
            else:
                ax = sns.boxplot(x=x_variable, y=y_variable, data=df, palette=config.get('color_palette', 'viridis'))
        
        else:
            return jsonify({
                'success': False,
                'message': f'不支持的图表类型: {chart_type}'
            }), 400
        
        # 设置标题和标签
        plt.title(config.get('title', ''))
        plt.xlabel(config.get('x_label', x_variable))
        plt.ylabel(config.get('y_label', y_variable if y_variable else ''))
        
        # 设置网格
        if config.get('grid', True):
            plt.grid(True, alpha=0.3)
        
        # 设置图例
        if config.get('legend', True) and group_variable:
            plt.legend(title=group_variable, loc=config.get('legend_position', 'best'))
        
        # 旋转X轴标签
        if config.get('rotation'):
            plt.xticks(rotation=config.get('rotation'))
        
        # 添加注释
        annotations = config.get('annotations', [])
        for annotation in annotations:
            plt.annotate(
                annotation.get('text', ''),
                xy=(annotation.get('x'), annotation.get('y')),
                color=annotation.get('color', 'red'),
                arrowprops=dict(facecolor=annotation.get('color', 'red'), shrink=0.05)
            )
        
        # 保存图表为PNG格式的base64字符串
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=config.get('dpi', 100))
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        plt.close()
        
        # 构建响应
        response = {
            'success': True,
            'message': '图表生成成功',
            'chart_data': {
                'type': chart_type,
                'image': f'data:image/png;base64,{image_base64}',
                'config': config
            }
        }
        
        return jsonify(response)
        
    except Exception as e:
        current_app.logger.error(f"高级图表配置API错误: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'处理请求时出错: {str(e)}'
        }), 500

# 交互式数据探索API
@advanced_viz_bp.route('/api/visualization/interactive_explore', methods=['POST', 'OPTIONS'])
@login_required
def interactive_explore():
    """交互式数据探索API
    
    允许用户通过交互方式探索数据，动态生成可视化
    
    请求体格式:
    {
        "dataset_id": 1,                     // 数据集ID
        "exploration_type": "distribution",  // 探索类型: distribution, correlation, pairplot, cluster
        "variables": ["age", "bmi", "gender"],  // 要探索的变量列表
        "config": {                          // 探索配置
            "n_bins": 20,                    // 直方图的箱数（仅用于distribution类型）
            "correlation_method": "pearson", // 相关性方法（仅用于correlation类型）
            "cluster_n": 3,                  // 聚类数量（仅用于cluster类型）
            "color_palette": "viridis",      // 颜色方案
            "figure_size": [10, 6],          // 图表尺寸 [宽, 高]
            "dpi": 100,                      // 分辨率
            "style": "whitegrid"             // 图表样式
        }
    }
    
    Returns:
        包含探索结果的JSON响应
    """
    if request.method == 'OPTIONS':
        response = current_app.make_default_options_response()
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST')
        return response
        
    try:
        # 获取请求数据
        request_data = request.get_json()
        if not request_data:
            return jsonify({
                'success': False,
                'message': '请求数据为空'
            }), 400
            
        # 提取参数
        dataset_id = request_data.get('dataset_id')
        exploration_type = request_data.get('exploration_type')
        variables = request_data.get('variables', [])
        config = request_data.get('config', {})
        
        # 验证必要参数
        if not dataset_id:
            return jsonify({
                'success': False,
                'message': '未指定数据集ID'
            }), 400
        
        if not exploration_type:
            return jsonify({
                'success': False,
                'message': '未指定探索类型'
            }), 400
        
        if not variables or len(variables) == 0:
            return jsonify({
                'success': False,
                'message': '未指定要探索的变量'
            }), 400
        
        # 根据探索类型验证其他必要参数
        if exploration_type == 'correlation' and len(variables) < 2:
            return jsonify({
                'success': False,
                'message': '相关性分析至少需要两个变量'
            }), 400
        
        # 检查数据集是否存在，并验证访问权限
        from app.models.dataset import DataSet
        from app.models.dataset_entry import DatasetEntry
        
        dataset = DataSet.query.get(dataset_id)
        if not dataset:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})不存在'
            }), 404
        
        # 检查用户是否有权限访问此数据集
        if dataset.privacy_level and dataset.privacy_level != 'public':
            if dataset.created_by != current_user.id and current_user.role != 'admin':
                if not dataset.is_shared_with(current_user.id):
                    return jsonify({
                        'success': False, 
                        'message': f'没有权限访问数据集(ID={dataset_id})'
                    }), 403
        
        # 获取数据集条目
        entries = DatasetEntry.query.filter_by(dataset_id=dataset_id).all()
        if not entries:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})没有数据条目'
            }), 404
        
        # 提取数据
        data = {}
        for var in variables:
            data[var] = []
        
        # 从所有条目中提取所需变量的值
        for entry in entries:
            try:
                entry_data = json.loads(entry.data)
                
                for var in variables:
                    if var in entry_data:
                        value = entry_data[var]
                        # 尝试将字符串转换为数值
                        if isinstance(value, str):
                            try:
                                if '.' in value:
                                    value = float(value)
                                else:
                                    value = int(value)
                            except (ValueError, TypeError):
                                pass
                        data[var].append(value)
                    else:
                        data[var].append(None)  # 缺失值
            except Exception as e:
                current_app.logger.error(f"处理条目数据时出错: {str(e)}")
                continue
        
        # 转换为Pandas DataFrame
        df = pd.DataFrame(data)
        
        # 去除包含缺失值的行
        df = df.dropna()
        
        if len(df) == 0:
            return jsonify({
                'success': False,
                'message': '处理后的数据为空，无法进行探索'
            }), 400
        
        # 应用图表配置
        plt.figure(figsize=tuple(config.get('figure_size', [10, 6])))
        sns.set_style(config.get('style', 'whitegrid'))
        
        # 根据探索类型生成可视化
        if exploration_type == 'distribution':
            # 分析变量分布
            fig, axes = plt.subplots(len(variables), 1, figsize=tuple(config.get('figure_size', [10, 6 * len(variables)])))
            
            # 如果只有一个变量，axes不会是数组
            if len(variables) == 1:
                axes = [axes]
            
            for i, var in enumerate(variables):
                # 检查变量类型
                if pd.api.types.is_numeric_dtype(df[var]):
                    # 数值型变量使用直方图
                    sns.histplot(df[var], bins=config.get('n_bins', 20), kde=True, ax=axes[i])
                    
                    # 添加描述性统计
                    mean = df[var].mean()
                    median = df[var].median()
                    std = df[var].std()
                    axes[i].axvline(mean, color='r', linestyle='--', label=f'均值: {mean:.2f}')
                    axes[i].axvline(median, color='g', linestyle='--', label=f'中位数: {median:.2f}')
                    axes[i].legend()
                    
                    # 设置标题和标签
                    axes[i].set_title(f'{var}的分布 (标准差: {std:.2f})')
                else:
                    # 分类变量使用计数图
                    sns.countplot(y=var, data=df, ax=axes[i], palette=config.get('color_palette', 'viridis'))
                    axes[i].set_title(f'{var}的分布')
            
            plt.tight_layout()
            
            # 计算描述性统计
            summary_stats = {}
            for var in variables:
                if pd.api.types.is_numeric_dtype(df[var]):
                    summary_stats[var] = {
                        'mean': float(df[var].mean()),
                        'median': float(df[var].median()),
                        'std': float(df[var].std()),
                        'min': float(df[var].min()),
                        'max': float(df[var].max()),
                        'q1': float(df[var].quantile(0.25)),
                        'q3': float(df[var].quantile(0.75))
                    }
                else:
                    # 分类变量的统计
                    value_counts = df[var].value_counts().to_dict()
                    summary_stats[var] = {
                        'categories': list(value_counts.keys()),
                        'counts': list(value_counts.values()),
                        'mode': df[var].mode()[0],
                        'unique_count': len(value_counts)
                    }
            
        elif exploration_type == 'correlation':
            # 筛选数值型变量
            numeric_vars = [var for var in variables if pd.api.types.is_numeric_dtype(df[var])]
            
            if len(numeric_vars) < 2:
                return jsonify({
                    'success': False,
                    'message': '相关性分析至少需要两个数值型变量'
                }), 400
            
            # 计算相关性矩阵
            corr_method = config.get('correlation_method', 'pearson')
            corr_matrix = df[numeric_vars].corr(method=corr_method)
            
            # 绘制热图
            plt.figure(figsize=tuple(config.get('figure_size', [10, 8])))
            mask = np.triu(np.ones_like(corr_matrix, dtype=bool))  # 创建上三角掩码
            ax = sns.heatmap(
                corr_matrix, 
                annot=True, 
                mask=mask,
                cmap=config.get('color_palette', 'viridis'),
                vmin=-1, vmax=1, 
                center=0,
                square=True, 
                linewidths=.5
            )
            plt.title(f'变量相关性矩阵 (方法: {corr_method})')
            
            # 保存相关性矩阵数据
            summary_stats = {
                'correlation_matrix': corr_matrix.to_dict(),
                'correlation_method': corr_method
            }
            
        elif exploration_type == 'pairplot':
            # 绘制配对图
            g = sns.pairplot(df[variables], palette=config.get('color_palette', 'viridis'))
            plt.suptitle('变量配对关系图', y=1.02)
            
            # 计算描述性统计
            summary_stats = {}
            for var in variables:
                if pd.api.types.is_numeric_dtype(df[var]):
                    summary_stats[var] = {
                        'mean': float(df[var].mean()),
                        'median': float(df[var].median()),
                        'std': float(df[var].std())
                    }
            
        elif exploration_type == 'cluster':
            # 聚类分析
            from sklearn.cluster import KMeans
            from sklearn.preprocessing import StandardScaler
            
            # 筛选数值型变量
            numeric_vars = [var for var in variables if pd.api.types.is_numeric_dtype(df[var])]
            
            if len(numeric_vars) < 2:
                return jsonify({
                    'success': False,
                    'message': '聚类分析至少需要两个数值型变量'
                }), 400
            
            # 标准化数据
            scaler = StandardScaler()
            scaled_data = scaler.fit_transform(df[numeric_vars])
            
            # 执行K-means聚类
            n_clusters = config.get('cluster_n', 3)
            kmeans = KMeans(n_clusters=n_clusters, random_state=42)
            df['cluster'] = kmeans.fit_predict(scaled_data)
            
            # 绘制聚类结果（选择前两个变量可视化）
            plt.figure(figsize=tuple(config.get('figure_size', [10, 8])))
            sns.scatterplot(
                x=numeric_vars[0],
                y=numeric_vars[1],
                hue='cluster',
                palette=config.get('color_palette', 'viridis'),
                data=df
            )
            plt.title(f'K-means聚类结果 (k={n_clusters})')
            
            # 计算每个聚类的统计信息
            cluster_stats = {}
            for i in range(n_clusters):
                cluster_df = df[df['cluster'] == i]
                cluster_stats[f'cluster_{i}'] = {
                    'size': len(cluster_df),
                    'percentage': float(len(cluster_df) / len(df) * 100),
                    'centroid': {var: float(cluster_df[var].mean()) for var in numeric_vars}
                }
            
            summary_stats = {
                'n_clusters': n_clusters,
                'cluster_stats': cluster_stats,
                'cluster_counts': df['cluster'].value_counts().to_dict()
            }
            
        else:
            return jsonify({
                'success': False,
                'message': f'不支持的探索类型: {exploration_type}'
            }), 400
        
        # 保存图表为PNG格式的base64字符串
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=config.get('dpi', 100))
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        plt.close()
        
        # 构建响应
        response = {
            'success': True,
            'message': '数据探索成功',
            'exploration_result': {
                'type': exploration_type,
                'image': f'data:image/png;base64,{image_base64}',
                'summary_stats': summary_stats,
                'config': config
            }
        }
        
        return jsonify(response)
        
    except Exception as e:
        current_app.logger.error(f"交互式数据探索API错误: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'处理请求时出错: {str(e)}'
        }), 500

# 自动报告生成API
@advanced_viz_bp.route('/api/visualization/generate_report', methods=['POST', 'OPTIONS'])
@login_required
def generate_report():
    """自动报告生成API
    
    根据分析结果自动生成报告，支持多种格式导出
    
    请求体格式:
    {
        "dataset_id": 1,                     // 数据集ID
        "report_type": "comprehensive",      // 报告类型: basic, comprehensive, custom
        "variables": ["age", "bmi", "gender", "smoking", "outcome"],  // 要包含在报告中的变量
        "outcome_variable": "outcome",       // 结局变量（可选）
        "format": "pdf",                     // 报告格式: pdf, docx, html
        "include_sections": {                // 要包含的报告部分
            "summary": true,                 // 数据摘要
            "descriptive": true,             // 描述性统计
            "distributions": true,           // 变量分布
            "correlations": true,            // 相关性分析
            "group_comparisons": true,       // 组间比较
            "predictions": false             // 预测分析
        },
        "title": "临床数据分析报告",          // 报告标题
        "author": "医学研究团队",            // 报告作者
        "config": {                          // 报告配置
            "color_theme": "professional",   // 报告颜色主题
            "logo_path": null,               // 徽标路径（可选）
            "page_size": "A4",               // 页面大小
            "include_code": false            // 是否包含代码
        }
    }
    
    Returns:
        包含报告文件URL的JSON响应或直接返回报告文件
    """
    if request.method == 'OPTIONS':
        response = current_app.make_default_options_response()
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST')
        return response
        
    try:
        # 获取请求数据
        request_data = request.get_json()
        if not request_data:
            return jsonify({
                'success': False,
                'message': '请求数据为空'
            }), 400
            
        # 提取参数
        dataset_id = request_data.get('dataset_id')
        report_type = request_data.get('report_type', 'comprehensive')
        variables = request_data.get('variables', [])
        outcome_variable = request_data.get('outcome_variable')
        report_format = request_data.get('format', 'pdf')
        include_sections = request_data.get('include_sections', {})
        title = request_data.get('title', '数据分析报告')
        author = request_data.get('author', '系统生成')
        config = request_data.get('config', {})
        
        # 验证必要参数
        if not dataset_id:
            return jsonify({
                'success': False,
                'message': '未指定数据集ID'
            }), 400
        
        if not variables or len(variables) == 0:
            return jsonify({
                'success': False,
                'message': '未指定要分析的变量'
            }), 400
        
        # 验证报告格式
        valid_formats = ['pdf', 'docx', 'html']
        if report_format not in valid_formats:
            return jsonify({
                'success': False,
                'message': f'不支持的报告格式: {report_format}，支持的格式有: {", ".join(valid_formats)}'
            }), 400
        
        # 检查数据集是否存在，并验证访问权限
        from app.models.dataset import DataSet
        from app.models.dataset_entry import DatasetEntry
        
        dataset = DataSet.query.get(dataset_id)
        if not dataset:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})不存在'
            }), 404
        
        # 检查用户是否有权限访问此数据集
        if dataset.privacy_level and dataset.privacy_level != 'public':
            if dataset.created_by != current_user.id and current_user.role != 'admin':
                if not dataset.is_shared_with(current_user.id):
                    return jsonify({
                        'success': False, 
                        'message': f'没有权限访问数据集(ID={dataset_id})'
                    }), 403
        
        # 获取数据集条目
        entries = DatasetEntry.query.filter_by(dataset_id=dataset_id).all()
        if not entries:
            return jsonify({
                'success': False,
                'message': f'数据集(ID={dataset_id})没有数据条目'
            }), 404
        
        # 提取数据
        data = {}
        for var in variables:
            data[var] = []
        
        # 从所有条目中提取所需变量的值
        for entry in entries:
            try:
                entry_data = json.loads(entry.data)
                
                for var in variables:
                    if var in entry_data:
                        value = entry_data[var]
                        # 尝试将字符串转换为数值
                        if isinstance(value, str):
                            try:
                                if '.' in value:
                                    value = float(value)
                                else:
                                    value = int(value)
                            except (ValueError, TypeError):
                                pass
                        data[var].append(value)
                    else:
                        data[var].append(None)  # 缺失值
            except Exception as e:
                current_app.logger.error(f"处理条目数据时出错: {str(e)}")
                continue
        
        # 转换为Pandas DataFrame
        df = pd.DataFrame(data)
        
        # 去除包含缺失值的行
        df = df.dropna()
        
        if len(df) == 0:
            return jsonify({
                'success': False,
                'message': '处理后的数据为空，无法生成报告'
            }), 400
            
        # 根据报告格式生成报告
        if report_format == 'pdf':
            report_path = generate_pdf_report(df, dataset, report_type, include_sections, title, author, config)
        elif report_format == 'docx':
            report_path = generate_docx_report(df, dataset, report_type, include_sections, title, author, config)
        elif report_format == 'html':
            report_path = generate_html_report(df, dataset, report_type, include_sections, title, author, config)
        else:
            return jsonify({
                'success': False,
                'message': f'不支持的报告格式: {report_format}'
            }), 400
        
        # 返回报告文件
        return send_file(
            report_path,
            as_attachment=True,
            download_name=f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{report_format}"
        )
        
    except Exception as e:
        current_app.logger.error(f"自动报告生成API错误: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'处理请求时出错: {str(e)}'
        }), 500

def generate_pdf_report(df, dataset, report_type, include_sections, title, author, config):
    """生成PDF格式的数据分析报告
    
    Args:
        df: 数据DataFrame
        dataset: 数据集对象
        report_type: 报告类型
        include_sections: 要包含的报告部分
        title: 报告标题
        author: 报告作者
        config: 报告配置
        
    Returns:
        生成的PDF报告文件路径
    """
    # 创建临时文件
    temp_dir = tempfile.mkdtemp()
    report_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    report_path = os.path.join(temp_dir, report_filename)
    
    # 设置页面大小
    page_size = A4
    if config.get('page_size') == 'letter':
        page_size = letter
    
    # 创建PDF文档
    doc = SimpleDocTemplate(
        report_path,
        pagesize=page_size,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    heading1_style = ParagraphStyle(
        'Heading1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12
    )
    
    heading2_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=10
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8
    )
    
    # 创建文档内容
    content = []
    
    # 添加标题
    content.append(Paragraph(title, title_style))
    content.append(Spacer(1, 0.25*inch))
    
    # 添加作者和日期
    content.append(Paragraph(f"作者: {author}", normal_style))
    content.append(Paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    content.append(Spacer(1, 0.5*inch))
    
    # 添加数据集信息
    content.append(Paragraph("数据集信息", heading1_style))
    content.append(Paragraph(f"数据集名称: {dataset.name}", normal_style))
    content.append(Paragraph(f"数据集描述: {dataset.description or '无'}", normal_style))
    content.append(Paragraph(f"样本数量: {len(df)}", normal_style))
    content.append(Paragraph(f"变量数量: {len(df.columns)}", normal_style))
    content.append(Spacer(1, 0.25*inch))
    
    # 添加数据摘要
    if include_sections.get('summary', True):
        content.append(Paragraph("数据摘要", heading1_style))
        
        # 创建变量表格
        var_data = [["变量名", "类型", "非空值数", "缺失率", "唯一值数"]]
        for col in df.columns:
            var_type = str(df[col].dtype)
            non_null_count = df[col].count()
            missing_rate = f"{(1 - non_null_count / len(df)) * 100:.2f}%"
            unique_count = df[col].nunique()
            var_data.append([col, var_type, str(non_null_count), missing_rate, str(unique_count)])
        
        var_table = Table(var_data, colWidths=[2*inch, 1.5*inch, 1*inch, 1*inch, 1*inch])
        var_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        content.append(var_table)
        content.append(Spacer(1, 0.25*inch))
    
    # 添加描述性统计
    if include_sections.get('descriptive', True):
        content.append(Paragraph("描述性统计", heading1_style))
        
        # 筛选数值型变量
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        if numeric_cols:
            # 计算描述性统计
            desc_stats = df[numeric_cols].describe().round(2)
            
            # 创建描述性统计表格
            desc_data = [["统计量"] + numeric_cols]
            for stat in desc_stats.index:
                row = [stat]
                for col in numeric_cols:
                    row.append(str(desc_stats.loc[stat, col]))
                desc_data.append(row)
            
            desc_table = Table(desc_data, colWidths=[1.2*inch] + [1.2*inch] * len(numeric_cols))
            desc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            content.append(desc_table)
            content.append(Spacer(1, 0.25*inch))
        else:
            content.append(Paragraph("没有可用于描述性统计的数值型变量", normal_style))
            content.append(Spacer(1, 0.25*inch))
        
        # 分类变量的频率统计
        cat_cols = df.select_dtypes(exclude=['number']).columns.tolist()
        if cat_cols:
            content.append(Paragraph("分类变量频率统计", heading2_style))
            
            for col in cat_cols:
                content.append(Paragraph(f"变量: {col}", normal_style))
                
                # 计算频率
                freq = df[col].value_counts().reset_index()
                freq.columns = [col, '频数']
                freq['频率'] = (freq['频数'] / freq['频数'].sum() * 100).round(2).astype(str) + '%'
                
                # 创建频率表格
                freq_data = [[col, '频数', '频率']]
                for _, row in freq.iterrows():
                    freq_data.append([str(row[col]), str(row['频数']), row['频率']])
                
                freq_table = Table(freq_data, colWidths=[2*inch, 1*inch, 1*inch])
                freq_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                content.append(freq_table)
                content.append(Spacer(1, 0.25*inch))
    
    # 添加变量分布
    if include_sections.get('distributions', True) and len(df.columns) > 0:
        content.append(Paragraph("变量分布", heading1_style))
        
        # 为每个数值型变量生成直方图
        for col in numeric_cols:
            plt.figure(figsize=(6, 4))
            sns.histplot(df[col], kde=True)
            plt.title(f"{col}的分布")
            plt.xlabel(col)
            plt.ylabel("频数")
            
            # 保存图表为临时文件
            img_path = os.path.join(temp_dir, f"{col}_hist.png")
            plt.savefig(img_path, dpi=100)
            plt.close()
            
            # 添加图表到报告
            content.append(Paragraph(f"{col}的分布", heading2_style))
            content.append(Image(img_path, width=5*inch, height=3*inch))
            content.append(Spacer(1, 0.25*inch))
    
    # 添加相关性分析
    if include_sections.get('correlations', True) and len(numeric_cols) >= 2:
        content.append(Paragraph("相关性分析", heading1_style))
        
        # 计算相关性矩阵
        corr_matrix = df[numeric_cols].corr()
        
        # 生成热图
        plt.figure(figsize=(8, 6))
        sns.heatmap(corr_matrix, annot=True, cmap='viridis', vmin=-1, vmax=1, center=0)
        plt.title("变量相关性热图")
        
        # 保存热图为临时文件
        corr_img_path = os.path.join(temp_dir, "correlation_heatmap.png")
        plt.savefig(corr_img_path, dpi=100)
        plt.close()
        
        # 添加热图到报告
        content.append(Paragraph("变量相关性热图", heading2_style))
        content.append(Image(corr_img_path, width=6*inch, height=4.5*inch))
        content.append(Spacer(1, 0.25*inch))
        
        # 添加相关性表格
        content.append(Paragraph("相关系数表", heading2_style))
        
        # 创建相关性表格
        corr_data = [["变量"] + numeric_cols]
        for col in numeric_cols:
            row = [col]
            for col2 in numeric_cols:
                row.append(f"{corr_matrix.loc[col, col2]:.2f}")
            corr_data.append(row)
        
        corr_table = Table(corr_data, colWidths=[1.2*inch] + [1*inch] * len(numeric_cols))
        corr_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        content.append(corr_table)
        content.append(Spacer(1, 0.25*inch))
    
    # 添加组间比较（如果有结局变量）
    if include_sections.get('group_comparisons', True) and 'outcome_variable' in locals() and outcome_variable in df.columns:
        content.append(Paragraph("组间比较", heading1_style))
        
        # 检查结局变量类型
        if df[outcome_variable].nunique() <= 5:  # 假设是分类变量
            # 对每个数值型变量进行箱线图比较
            for col in numeric_cols:
                if col != outcome_variable:
                    plt.figure(figsize=(6, 4))
                    sns.boxplot(x=outcome_variable, y=col, data=df)
                    plt.title(f"按{outcome_variable}分组的{col}箱线图")
                    
                    # 保存图表为临时文件
                    box_img_path = os.path.join(temp_dir, f"{col}_by_{outcome_variable}_box.png")
                    plt.savefig(box_img_path, dpi=100)
                    plt.close()
                    
                    # 添加图表到报告
                    content.append(Paragraph(f"按{outcome_variable}分组的{col}箱线图", heading2_style))
                    content.append(Image(box_img_path, width=5*inch, height=3*inch))
                    content.append(Spacer(1, 0.25*inch))
            
            # 添加分组统计表格
            content.append(Paragraph("分组统计表", heading2_style))
            
            # 计算每个组的统计量
            group_stats = {}
            for group in df[outcome_variable].unique():
                group_df = df[df[outcome_variable] == group]
                group_stats[group] = {}
                for col in numeric_cols:
                    if col != outcome_variable:
                        group_stats[group][col] = {
                            'mean': group_df[col].mean(),
                            'std': group_df[col].std(),
                            'median': group_df[col].median(),
                            'min': group_df[col].min(),
                            'max': group_df[col].max(),
                            'count': group_df[col].count()
                        }
            
            # 为每个数值变量创建分组统计表
            for col in numeric_cols:
                if col != outcome_variable:
                    content.append(Paragraph(f"{col}的分组统计", normal_style))
                    
                    # 创建表格
                    stat_data = [["组别", "样本数", "均值", "标准差", "中位数", "最小值", "最大值"]]
                    for group in sorted(group_stats.keys()):
                        stats = group_stats[group][col]
                        stat_data.append([
                            str(group),
                            str(int(stats['count'])),
                            f"{stats['mean']:.2f}",
                            f"{stats['std']:.2f}",
                            f"{stats['median']:.2f}",
                            f"{stats['min']:.2f}",
                            f"{stats['max']:.2f}"
                        ])
                    
                    stat_table = Table(stat_data, colWidths=[0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch])
                    stat_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    content.append(stat_table)
                    content.append(Spacer(1, 0.25*inch))
    
    # 构建PDF文档
    doc.build(content)
    
    return report_path

def generate_docx_report(df, dataset, report_type, include_sections, title, author, config):
    """生成DOCX格式的数据分析报告
    
    Args:
        df: 数据DataFrame
        dataset: 数据集对象
        report_type: 报告类型
        include_sections: 要包含的报告部分
        title: 报告标题
        author: 报告作者
        config: 报告配置
        
    Returns:
        生成的DOCX报告文件路径
    """
    # 创建临时文件
    temp_dir = tempfile.mkdtemp()
    report_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    report_path = os.path.join(temp_dir, report_filename)
    
    # 创建文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # 添加标题
    doc.add_heading(title, level=0)
    
    # 添加作者和日期
    doc.add_paragraph(f"作者: {author}")
    doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()  # 空行
    
    # 添加数据集信息
    doc.add_heading("数据集信息", level=1)
    doc.add_paragraph(f"数据集名称: {dataset.name}")
    doc.add_paragraph(f"数据集描述: {dataset.description or '无'}")
    doc.add_paragraph(f"样本数量: {len(df)}")
    doc.add_paragraph(f"变量数量: {len(df.columns)}")
    doc.add_paragraph()  # 空行
    
    # 添加数据摘要
    if include_sections.get('summary', True):
        doc.add_heading("数据摘要", level=1)
        
        # 创建变量表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "变量名"
        header_cells[1].text = "类型"
        header_cells[2].text = "非空值数"
        header_cells[3].text = "缺失率"
        header_cells[4].text = "唯一值数"
        
        # 设置表头格式
        for cell in header_cells:
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True
        
        # 添加变量数据
        for col in df.columns:
            row_cells = table.add_row().cells
            var_type = str(df[col].dtype)
            non_null_count = df[col].count()
            missing_rate = f"{(1 - non_null_count / len(df)) * 100:.2f}%"
            unique_count = df[col].nunique()
            
            row_cells[0].text = col
            row_cells[1].text = var_type
            row_cells[2].text = str(non_null_count)
            row_cells[3].text = missing_rate
            row_cells[4].text = str(unique_count)
            
            # 设置单元格文本居中
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
    
    # 添加描述性统计
    if include_sections.get('descriptive', True):
        doc.add_heading("描述性统计", level=1)
        
        # 筛选数值型变量
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        if numeric_cols:
            # 计算描述性统计
            desc_stats = df[numeric_cols].describe().round(2)
            
            # 创建描述性统计表格
            table = doc.add_table(rows=1, cols=len(numeric_cols) + 1)
            table.style = 'Table Grid'
            
            # 设置表头
            header_cells = table.rows[0].cells
            header_cells[0].text = "统计量"
            for i, col in enumerate(numeric_cols):
                header_cells[i + 1].text = col
            
            # 设置表头格式
            for cell in header_cells:
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell_para.runs:
                    run.bold = True
            
            # 添加统计数据
            for stat in desc_stats.index:
                row_cells = table.add_row().cells
                row_cells[0].text = stat
                for i, col in enumerate(numeric_cols):
                    row_cells[i + 1].text = str(desc_stats.loc[stat, col])
                
                # 设置单元格文本居中
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # 空行
        else:
            doc.add_paragraph("没有可用于描述性统计的数值型变量")
            doc.add_paragraph()  # 空行
        
        # 分类变量的频率统计
        cat_cols = df.select_dtypes(exclude=['number']).columns.tolist()
        if cat_cols:
            doc.add_heading("分类变量频率统计", level=2)
            
            for col in cat_cols:
                doc.add_paragraph(f"变量: {col}")
                
                # 计算频率
                freq = df[col].value_counts().reset_index()
                freq.columns = [col, '频数']
                freq['频率'] = (freq['频数'] / freq['频数'].sum() * 100).round(2).astype(str) + '%'
                
                # 创建频率表格
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # 设置表头
                header_cells = table.rows[0].cells
                header_cells[0].text = col
                header_cells[1].text = "频数"
                header_cells[2].text = "频率"
                
                # 设置表头格式
                for cell in header_cells:
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell_para.runs:
                        run.bold = True
                
                # 添加频率数据
                for _, row in freq.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row[col])
                    row_cells[1].text = str(row['频数'])
                    row_cells[2].text = row['频率']
                    
                    # 设置单元格文本居中
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # 空行
    
    # 添加变量分布
    if include_sections.get('distributions', True) and len(df.columns) > 0:
        doc.add_heading("变量分布", level=1)
        
        # 为每个数值型变量生成直方图
        for col in numeric_cols:
            plt.figure(figsize=(6, 4))
            sns.histplot(df[col], kde=True)
            plt.title(f"{col}的分布")
            plt.xlabel(col)
            plt.ylabel("频数")
            
            # 保存图表为临时文件
            img_path = os.path.join(temp_dir, f"{col}_hist.png")
            plt.savefig(img_path, dpi=100)
            plt.close()
            
            # 添加图表标题和图片
            doc.add_heading(f"{col}的分布", level=2)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()  # 空行
    
    # 添加相关性分析
    if include_sections.get('correlations', True) and len(numeric_cols) >= 2:
        doc.add_heading("相关性分析", level=1)
        
        # 计算相关性矩阵
        corr_matrix = df[numeric_cols].corr()
        
        # 生成热图
        plt.figure(figsize=(8, 6))
        sns.heatmap(corr_matrix, annot=True, cmap='viridis', vmin=-1, vmax=1, center=0)
        plt.title("变量相关性热图")
        
        # 保存热图为临时文件
        corr_img_path = os.path.join(temp_dir, "correlation_heatmap.png")
        plt.savefig(corr_img_path, dpi=100)
        plt.close()
        
        # 添加热图标题和图片
        doc.add_heading("变量相关性热图", level=2)
        doc.add_picture(corr_img_path, width=Inches(6))
        doc.add_paragraph()  # 空行
        
        # 添加相关性表格
        doc.add_heading("相关系数表", level=2)
        
        # 创建相关性表格
        table = doc.add_table(rows=1, cols=len(numeric_cols) + 1)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "变量"
        for i, col in enumerate(numeric_cols):
            header_cells[i + 1].text = col
        
        # 设置表头格式
        for cell in header_cells:
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True
        
        # 添加相关性数据
        for i, col in enumerate(numeric_cols):
            row_cells = table.add_row().cells
            row_cells[0].text = col
            for j, col2 in enumerate(numeric_cols):
                row_cells[j + 1].text = f"{corr_matrix.loc[col, col2]:.2f}"
            
            # 设置单元格文本居中
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
    
    # 保存文档
    doc.save(report_path)
    
    return report_path

def generate_html_report(df, dataset, report_type, include_sections, title, author, config):
    """生成HTML格式的数据分析报告
    
    Args:
        df: 数据DataFrame
        dataset: 数据集对象
        report_type: 报告类型
        include_sections: 要包含的报告部分
        title: 报告标题
        author: 报告作者
        config: 报告配置
        
    Returns:
        生成的HTML报告文件路径
    """
    # 创建临时文件
    temp_dir = tempfile.mkdtemp()
    report_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    report_path = os.path.join(temp_dir, report_filename)
    
    # 设置颜色主题
    color_theme = config.get('color_theme', 'professional')
    if color_theme == 'professional':
        primary_color = '#2c3e50'
        secondary_color = '#3498db'
        background_color = '#f9f9f9'
        text_color = '#333333'
    elif color_theme == 'light':
        primary_color = '#4CAF50'
        secondary_color = '#8BC34A'
        background_color = '#ffffff'
        text_color = '#212121'
    elif color_theme == 'dark':
        primary_color = '#303F9F'
        secondary_color = '#3F51B5'
        background_color = '#212121'
        text_color = '#FFFFFF'
    else:
        primary_color = '#2c3e50'
        secondary_color = '#3498db'
        background_color = '#f9f9f9'
        text_color = '#333333'
    
    # 创建HTML头部
    html_content = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: {text_color};
                background-color: {background_color};
                margin: 0;
                padding: 20px;
            }}
            .container {{
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
                background-color: #fff;
                box-shadow: 0 0 10px rgba(0,0,0,0.1);
            }}
            h1, h2, h3 {{
                color: {primary_color};
                margin-top: 20px;
            }}
            h1 {{
                text-align: center;
                padding-bottom: 10px;
                border-bottom: 2px solid {secondary_color};
            }}
            h2 {{
                border-bottom: 1px solid {secondary_color};
                padding-bottom: 5px;
            }}
            .meta-info {{
                margin-bottom: 30px;
                text-align: center;
                color: #666;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th, td {{
                padding: 12px 15px;
                text-align: center;
                border: 1px solid #ddd;
            }}
            th {{
                background-color: {secondary_color};
                color: white;
                font-weight: bold;
            }}
            tr:nth-child(even) {{
                background-color: #f2f2f2;
            }}
            .chart-container {{
                margin: 20px 0;
                text-align: center;
            }}
            .chart-container img {{
                max-width: 100%;
                height: auto;
                border: 1px solid #ddd;
            }}
            .footer {{
                margin-top: 50px;
                text-align: center;
                font-size: 0.9em;
                color: #666;
                border-top: 1px solid #ddd;
                padding-top: 20px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>{title}</h1>
            <div class="meta-info">
                <p>作者: {author} | 生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            </div>
    """
    
    # 添加数据集信息
    html_content += f"""
            <h2>数据集信息</h2>
            <p>数据集名称: {dataset.name}</p>
            <p>数据集描述: {dataset.description or '无'}</p>
            <p>样本数量: {len(df)}</p>
            <p>变量数量: {len(df.columns)}</p>
    """
    
    # 添加数据摘要
    if include_sections.get('summary', True):
        html_content += """
            <h2>数据摘要</h2>
            <table>
                <tr>
                    <th>变量名</th>
                    <th>类型</th>
                    <th>非空值数</th>
                    <th>缺失率</th>
                    <th>唯一值数</th>
                </tr>
        """
        
        for col in df.columns:
            var_type = str(df[col].dtype)
            non_null_count = df[col].count()
            missing_rate = f"{(1 - non_null_count / len(df)) * 100:.2f}%"
            unique_count = df[col].nunique()
            
            html_content += f"""
                <tr>
                    <td>{col}</td>
                    <td>{var_type}</td>
                    <td>{non_null_count}</td>
                    <td>{missing_rate}</td>
                    <td>{unique_count}</td>
                </tr>
            """
        
        html_content += """
            </table>
        """
    
    # 添加描述性统计
    if include_sections.get('descriptive', True):
        html_content += """
            <h2>描述性统计</h2>
        """
        
        # 筛选数值型变量
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        if numeric_cols:
            # 计算描述性统计
            desc_stats = df[numeric_cols].describe().round(2)
            
            # 创建描述性统计表格
            html_content += """
                <table>
                    <tr>
                        <th>统计量</th>
            """
            
            for col in numeric_cols:
                html_content += f"<th>{col}</th>"
            
            html_content += """
                    </tr>
            """
            
            for stat in desc_stats.index:
                html_content += f"""
                    <tr>
                        <td>{stat}</td>
                """
                
                for col in numeric_cols:
                    html_content += f"<td>{desc_stats.loc[stat, col]}</td>"
                
                html_content += """
                    </tr>
                """
            
            html_content += """
                </table>
            """
        else:
            html_content += """
                <p>没有可用于描述性统计的数值型变量</p>
            """
        
        # 分类变量的频率统计
        cat_cols = df.select_dtypes(exclude=['number']).columns.tolist()
        if cat_cols:
            html_content += """
                <h3>分类变量频率统计</h3>
            """
            
            for col in cat_cols:
                html_content += f"""
                    <h4>变量: {col}</h4>
                    <table>
                        <tr>
                            <th>{col}</th>
                            <th>频数</th>
                            <th>频率</th>
                        </tr>
                """
                
                # 计算频率
                freq = df[col].value_counts().reset_index()
                freq.columns = [col, '频数']
                freq['频率'] = (freq['频数'] / freq['频数'].sum() * 100).round(2).astype(str) + '%'
                
                for _, row in freq.iterrows():
                    html_content += f"""
                        <tr>
                            <td>{row[col]}</td>
                            <td>{row['频数']}</td>
                            <td>{row['频率']}</td>
                        </tr>
                    """
                
                html_content += """
                    </table>
                """
    
    # 添加变量分布
    if include_sections.get('distributions', True) and len(df.columns) > 0:
        html_content += """
            <h2>变量分布</h2>
        """
        
        # 为每个数值型变量生成直方图
        for col in numeric_cols:
            plt.figure(figsize=(8, 5))
            sns.histplot(df[col], kde=True)
            plt.title(f"{col}的分布")
            plt.xlabel(col)
            plt.ylabel("频数")
            
            # 保存图表为临时文件
            img_path = os.path.join(temp_dir, f"{col}_hist.png")
            plt.savefig(img_path, dpi=100)
            plt.close()
            
            # 将图片转换为base64编码，嵌入HTML
            with open(img_path, "rb") as img_file:
                img_base64 = base64.b64encode(img_file.read()).decode('utf-8')
            
            html_content += f"""
                <h3>{col}的分布</h3>
                <div class="chart-container">
                    <img src="data:image/png;base64,{img_base64}" alt="{col}的分布">
                </div>
            """
    
    # 添加相关性分析
    if include_sections.get('correlations', True) and len(numeric_cols) >= 2:
        html_content += """
            <h2>相关性分析</h2>
        """
        
        # 计算相关性矩阵
        corr_matrix = df[numeric_cols].corr()
        
        # 生成热图
        plt.figure(figsize=(10, 8))
        sns.heatmap(corr_matrix, annot=True, cmap='viridis', vmin=-1, vmax=1, center=0)
        plt.title("变量相关性热图")
        
        # 保存热图为临时文件
        corr_img_path = os.path.join(temp_dir, "correlation_heatmap.png")
        plt.savefig(corr_img_path, dpi=100)
        plt.close()
        
        # 将图片转换为base64编码，嵌入HTML
        with open(corr_img_path, "rb") as img_file:
            corr_img_base64 = base64.b64encode(img_file.read()).decode('utf-8')
        
        html_content += f"""
            <h3>变量相关性热图</h3>
            <div class="chart-container">
                <img src="data:image/png;base64,{corr_img_base64}" alt="变量相关性热图">
            </div>
            
            <h3>相关系数表</h3>
            <table>
                <tr>
                    <th>变量</th>
        """
        
        for col in numeric_cols:
            html_content += f"<th>{col}</th>"
        
        html_content += """
                </tr>
        """
        
        for col in numeric_cols:
            html_content += f"""
                <tr>
                    <td>{col}</td>
            """
            
            for col2 in numeric_cols:
                html_content += f"<td>{corr_matrix.loc[col, col2]:.2f}</td>"
            
            html_content += """
                </tr>
            """
        
        html_content += """
            </table>
        """
    
    # 添加组间比较（如果有结局变量）
    if include_sections.get('group_comparisons', True) and 'outcome_variable' in locals() and outcome_variable in df.columns:
        html_content += f"""
            <h2>组间比较</h2>
        """
        
        # 检查结局变量类型
        if df[outcome_variable].nunique() <= 5:  # 假设是分类变量
            # 对每个数值型变量进行箱线图比较
            for col in numeric_cols:
                if col != outcome_variable:
                    plt.figure(figsize=(8, 5))
                    sns.boxplot(x=outcome_variable, y=col, data=df)
                    plt.title(f"按{outcome_variable}分组的{col}箱线图")
                    
                    # 保存图表为临时文件
                    box_img_path = os.path.join(temp_dir, f"{col}_by_{outcome_variable}_box.png")
                    plt.savefig(box_img_path, dpi=100)
                    plt.close()
                    
                    # 将图片转换为base64编码，嵌入HTML
                    with open(box_img_path, "rb") as img_file:
                        box_img_base64 = base64.b64encode(img_file.read()).decode('utf-8')
                    
                    html_content += f"""
                        <h3>按{outcome_variable}分组的{col}箱线图</h3>
                        <div class="chart-container">
                            <img src="data:image/png;base64,{box_img_base64}" alt="按{outcome_variable}分组的{col}箱线图">
                        </div>
                    """
            
            # 添加分组统计表格
            html_content += """
                <h3>分组统计表</h3>
            """
            
            # 计算每个组的统计量
            group_stats = {}
            for group in df[outcome_variable].unique():
                group_df = df[df[outcome_variable] == group]
                group_stats[group] = {}
                for col in numeric_cols:
                    if col != outcome_variable:
                        group_stats[group][col] = {
                            'mean': group_df[col].mean(),
                            'std': group_df[col].std(),
                            'median': group_df[col].median(),
                            'min': group_df[col].min(),
                            'max': group_df[col].max(),
                            'count': group_df[col].count()
                        }
            
            # 为每个数值变量创建分组统计表
            for col in numeric_cols:
                if col != outcome_variable:
                    html_content += f"""
                        <h4>{col}的分组统计</h4>
                        <table>
                            <tr>
                                <th>组别</th>
                                <th>样本数</th>
                                <th>均值</th>
                                <th>标准差</th>
                                <th>中位数</th>
                                <th>最小值</th>
                                <th>最大值</th>
                            </tr>
                    """
                    
                    for group in sorted(group_stats.keys()):
                        stats = group_stats[group][col]
                        html_content += f"""
                            <tr>
                                <td>{group}</td>
                                <td>{int(stats['count'])}</td>
                                <td>{stats['mean']:.2f}</td>
                                <td>{stats['std']:.2f}</td>
                                <td>{stats['median']:.2f}</td>
                                <td>{stats['min']:.2f}</td>
                                <td>{stats['max']:.2f}</td>
                            </tr>
                        """
                    
                    html_content += """
                        </table>
                    """
    
    # 添加页脚
    html_content += f"""
            <div class="footer">
                <p>报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                <p>© {datetime.now().year} 自动报告生成系统</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    # 写入HTML文件
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return report_path 